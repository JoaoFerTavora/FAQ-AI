# LOGICA DA IA
import time
import io
import tiktoken as tik
import pypdf
import docx
from g4f.client import Client
from groq import Groq
from config import PROMPT_SISTEMA_SENAI, MODELO_DEFAULT_G4F, MODELO_DEFAULT_GROQ

def calcular_tokens(texto, modelo="gpt-4o-mini"):
    try:
        # Puxa o codificador oficial da OpenAI
        # configurado para o modelo
        codificador = tik.encoding_for_model(modelo)
    except KeyError:
        # Caso o modelo seja genérico,
        # usa o padrão do GPT-4
        codificador = tik.encoding_for_model("cl100k_base")
    
    # O método .encode() transformar o texto puro
    # em uma lista de números (IDs dos tokens)
    lista_de_tokens = codificador.encode(texto)

    # Retornar o tamanho dessa lista
    # que é a quantidade exata de tokens
    return len(lista_de_tokens)

# Função para ler PDF/DOCX
def extrair_texto_de_arquivo(uploaded_file):
    # Verifica a extensão do arquivo
    if uploaded_file is None:
        return ""
    if uploaded_file.name.endswith('.txt'):
        return uploaded_file.getvalue().decode('utf-8')
    ## Fazer o if com .docx
    elif uploaded_file.name.endswith('.docx'):
        doc = docx.Document(uploaded_file)
        return "\n".join([p.text for p in doc.paragraphs if p.text])
    # Fazer o if com .pdf
    elif uploaded_file.name.endswith('.pdf'):
        leitor_pdf = pypdf.PdfReader(uploaded_file)
        testoextraido = ""
        for pagina in leitor_pdf.pages:
            t = pagina.extract_text()
            if t:
                textoextraido +=t + "\n"
        return textoextraido
    return ""

def gerar_template_word_bytes():
    doc = docx.Document()
    doc.add_heading("Base de dados oficial - Cursos Senai")
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0) # reset memory pointer of buffer
    return buffer

def gerar_resposta_ia(provedor, mensagem, groq_key = None, temperatura = 0.2):
    inicio = time.time()

    if provedor == "GPT-4o Mini (Via G4F)":
        client = Client()
        resposta = client.chat.completions.create(
            model = MODELO_DEFAULT_G4F,
            messages = mensagem,
            temperature = temperatura
        )
        texto = resposta.choices[0].message.content
    elif provedor == "Llama 3.3 (Via Groq)":
        if not groq_key:
            raise ValueError("GROQ_API_KEY nâo configurada no secrets.toml")
        client_groq = Groq(api_key = groq_key)
        resposta = client_groq.chat.completions.create(
            model = MODELO_DEFAULT_GROQ,
            messages = mensagem,
            temperature = temperatura
        )
        texto = resposta.choices[0].message.content
    else:
        raise ValueError("Provedor inválido.")
    tempo = round(time.time() - inicio, 2)
    return texto, tempo
